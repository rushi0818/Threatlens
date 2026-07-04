import re

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


SUSPICIOUS_KEYWORDS = [
    "verify your account", "click here", "urgent", "suspended",
    "confirm your details", "update your payment", "account locked",
    "dear customer", "prize", "winner", "bank details", "otp"
]


def extract_urls_from_text(text):
    pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
    return re.findall(pattern, text)


def scan_doc(file_path):
    """
    Opens a .docx file, reads all paragraphs and tables,
    then checks for phishing indicators.
    """
    if not DOCX_AVAILABLE:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    try:
        doc       = Document(file_path)
        full_text = ""

        # Read all paragraphs
        for para in doc.paragraphs:
            full_text += para.text + "\n"

        # Read text inside tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += cell.text + "\n"

    except Exception as e:
        return {"error": f"Could not read Word file: {str(e)}"}

    if not full_text.strip():
        return {"error": "Word document appears to be empty"}

    text_lower = full_text.lower()
    reasons    = []
    score      = 0

    # Check URLs
    urls = extract_urls_from_text(full_text)
    suspicious_urls = []
    for url in urls:
        url_lower = url.lower()
        if any(tld in url_lower for tld in [".xyz", ".tk", ".ml", ".ga", ".cf"]):
            suspicious_urls.append(url)
            score += 3
        if re.search(r"(\d{1,3}\.){3}\d{1,3}", url):
            suspicious_urls.append(url)
            score += 3

    if suspicious_urls:
        reasons.append(f"Suspicious URLs in document: {suspicious_urls[0][:60]}")

    # Check keywords
    found_keywords = [kw for kw in SUSPICIOUS_KEYWORDS if kw in text_lower]
    if found_keywords:
        score += len(found_keywords) * 2
        reasons.append(f"Phishing phrases found: {', '.join(found_keywords[:3])}")

    # Check brand impersonation
    brands = ["paypal", "amazon", "apple", "microsoft", "google",
              "netflix", "bank", "ebay", "instagram", "facebook"]
    found_brands = [b for b in brands if b in text_lower]
    if found_brands:
        score += 2
        reasons.append(f"Brand impersonation detected: {', '.join(found_brands)}")

    # Check for macros warning - we can't run them but we warn
    if "macros" in text_lower or "enable content" in text_lower:
        score += 4
        reasons.append("Document asks to enable macros — very common in malware docs")

    # Verdict
    if score >= 5:
        verdict     = "PHISHING"
        is_phishing = True
        confidence  = min(97, 55 + score * 3)
    elif score >= 2:
        verdict     = "SUSPICIOUS"
        is_phishing = True
        confidence  = min(70, 40 + score * 5)
    else:
        verdict     = "LEGITIMATE"
        is_phishing = False
        confidence  = 85

    return {
        "verdict"      : verdict,
        "is_phishing"  : is_phishing,
        "confidence"   : confidence,
        "score"        : score,
        "reasons"      : reasons,
        "urls_found"   : urls,
        "text_preview" : full_text[:300].strip(),
        "scan_type"    : "Word Document"
    }